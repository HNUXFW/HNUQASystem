from typing import List, Dict
from pathlib import Path
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
from app.config import settings
import fitz
import spacy
import cv2
import pytesseract
from docx import Document

class DocumentProcessor:
    def __init__(self):
        self.model = SentenceTransformer(settings.EMBEDDING_MODEL)
        self.index = None
        self.documents: List[Dict] = []

    def process_documents(self, docs_dir: Path) -> None:
        """处理目录下的所有文本、PDF和Word文件"""
        self.documents = []
        nlp = spacy.load("zh_core_web_sm")
        print(f"进入该函数{docs_dir}")
        for file_path in docs_dir.glob("**/*"):
            try:
                print(f"开始处理文件: {file_path}")
                if file_path.suffix.lower() == ".txt":
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = {'text': f.read(), 'images': []}
                elif file_path.suffix.lower() == ".pdf":
                    content = self._extract_pdf_content(file_path)
                elif file_path.suffix.lower() in [".docx", ".doc"]:
                    if file_path.suffix.lower() == ".doc":
                        print(f"Warning: Old .doc format not fully supported: {file_path}")
                        continue
                    content = self._extract_docx_content(file_path)
                else:
                    continue

                # 处理文本内容
                if content['text']:
                    text_chunks = self._semantic_split(content['text'], nlp)
                    for i, chunk in enumerate(text_chunks):
                        self.documents.append({
                            "text": chunk,
                            "source": str(file_path),
                            "chunk_id": f"text_{i}",
                            "type": "text"
                        })
                
                # 处理图片中提取的文字内容
                for i, img in enumerate(content['images']):
                    if img['extracted_text']:
                        self.documents.append({
                            "text": img['extracted_text'],
                            "source": str(file_path),
                            "chunk_id": f"image_{i}",
                            "type": img['type'],
                            "page_num": img.get('page_num', 1)
                        })
            except Exception as e:
                print(f"Error processing file {file_path}: {e}")

    def _extract_pdf_content(self, pdf_path: Path) -> Dict:
        """提取PDF中的文本和图像内容"""
        doc = fitz.open(pdf_path)
        text_content = []
        image_content = []
        
        for page in doc:
            # 提取文本
            text = page.get_text()
            if text.strip():
                text_content.append(text)
            
            # 提取图像
            for img_index, img in enumerate(page.get_images()):
                xref = img[0]
                base_image = doc.extract_image(xref)
                
                # 处理图片内容
                image_result = self._process_image_content(
                    base_image["image"],
                    base_image["ext"]
                )
                
                # 只保存包含文字的图片信息
                if image_result['has_text']:
                    image_content.append({
                        'page_num': page.number,
                        'image_index': img_index,
                        'extracted_text': image_result['extracted_text'],
                        'type': image_result['type']
                    })
        
        return {
            'text': '\n'.join(text_content),
            'images': image_content
        }

    
    def build_index(self) -> None:
        """构建向量索引"""
        print("开始构建向量索引...")
        if not self.documents:
            raise ValueError("No documents processed yet")
            
        # 获取文档嵌入
        embeddings = self.model.encode([doc["text"] for doc in self.documents])
        
        # 构建FAISS索引
        dimension = embeddings.shape[1]
        self.index = faiss.IndexFlatL2(dimension)
        self.index.add(np.array(embeddings).astype('float32'))
        print("构建向量索引完成")
        
    def save_index(self, path: Path) -> None:
        """保存索引到文件"""
        if self.index is None:
            raise ValueError("No index built yet")
        path.parent.mkdir(parents=True, exist_ok=True)
        faiss.write_index(self.index, str(path))
        print(f"向量索引已保存到: {path}")
        
    def load_index(self, path: Path) -> None:
        """从文件加载索引"""
        if not path.exists():
            raise FileNotFoundError(f"Index file not found: {path}")
        self.index = faiss.read_index(str(path))
        
    def search(self, query: str, k: int = None) -> List[Dict]:
        """搜索最相关的文档片段"""
        if self.index is None:
            raise ValueError("No index built yet")
            
        k = k or settings.TOP_K
        query_vector = self.model.encode([query])
        distances, indices = self.index.search(np.array(query_vector).astype('float32'), k)
        
        results = []
        for idx in indices[0]:
            if idx != -1:  # FAISS可能返回-1表示无效结果
                results.append(self.documents[idx])
        return results 

    def _process_image_content(self, image_data: bytes, image_ext: str) -> Dict[str, any]:
        """处理图片内容，返回提取的文字和图片类型"""


        # 将图片数据转换为OpenCV格式
        image_array = np.frombuffer(image_data, np.uint8)
        img = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
        
        # 检测是否为流程图
        def is_flowchart(img):
            # 使用边缘检测和形状分析来判断是否为流程图
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            edges = cv2.Canny(gray, 50, 150)
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # 检查是否存在规则的几何形状（矩形、菱形等）
            for contour in contours:
                approx = cv2.approxPolyDP(contour, 0.04 * cv2.arcLength(contour, True), True)
                if len(approx) in [4, 6, 8]:  # 流程图常见的形状边数
                    return True
            return False

        # 图片预处理
        def preprocess_image(img):
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            denoised = cv2.fastNlMeansDenoising(gray)
            _, binary = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            return binary

        # 检测图片是否包含文字
        def has_text(img):
            preprocessed = preprocess_image(img)
            text = pytesseract.image_to_string(preprocessed, lang='chi_sim+eng')
            return len(text.strip()) > 0

        # 处理流程图
        def process_flowchart(img):
            preprocessed = preprocess_image(img)
            # 提取流程图中的文字
            text = pytesseract.image_to_string(preprocessed, lang='chi_sim+eng')
            #TODO：可以添加更多流程图专门的处理逻辑
            return text.strip()

        # 处理普通图片中的文字
        def process_regular_image(img):
            preprocessed = preprocess_image(img)
            text = pytesseract.image_to_string(preprocessed, lang='chi_sim+eng')
            return text.strip()

        result = {
            'has_text': False,
            'is_flowchart': False,
            'extracted_text': '',
            'type': 'image'
        }

        # 判断图片类型并处理
        if is_flowchart(img):
            #如果是流程图，那么识别流程图提取文字
            result['is_flowchart'] = True
            result['type'] = 'flowchart'
            extracted_text = process_flowchart(img)
            if extracted_text:
                result['has_text'] = True
                result['extracted_text'] = extracted_text
        elif has_text(img):
            #如果不是，那么判断是否包含文字，如果是，提取文字。
            result['has_text'] = True
            result['extracted_text'] = process_regular_image(img)

        return result 

    def _semantic_split(self, text: str, nlp) -> List[str]:
        """使用spaCy进行语义分块"""
        doc = nlp(text)
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sent in doc.sents:
            sent_text = sent.text.strip()
            # 跳过空句子
            if not sent_text:
                continue
            
            sent_length = len(sent_text)
            
            # 如果单个句子就超过了块大小，则需要进行分词切分
            if sent_length > settings.CHUNK_SIZE:
                # 如果当前chunk不为空，先保存
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                    current_chunk = []
                    current_length = 0
                
                # 对长句子进行分词切分
                words = [token.text for token in nlp(sent_text)]
                temp_chunk = []
                temp_length = 0
                
                for word in words:
                    if temp_length + len(word) < settings.CHUNK_SIZE:
                        temp_chunk.append(word)
                        temp_length += len(word)
                    else:
                        chunks.append(" ".join(temp_chunk))
                        temp_chunk = [word]
                        temp_length = len(word)
                
                if temp_chunk:
                    chunks.append(" ".join(temp_chunk))
                continue
            
            # 正常句子处理
            if current_length + sent_length < settings.CHUNK_SIZE:
                current_chunk.append(sent_text)
                current_length += sent_length
            else:
                chunks.append(" ".join(current_chunk))
                current_chunk = [sent_text]
                current_length = sent_length
        
        # 处理最后剩余的chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks 

    def _extract_docx_content(self, docx_path: Path) -> Dict:
        """提取Word文档中的文本和图片内容"""
        doc = Document(docx_path)
        text_content = []
        image_content = []
        
        # 提取文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        # 提取图片
        image_index = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    # 处理图片内容
                    image_result = self._process_image_content(image_data, rel.target_ref.split('.')[-1])
                    
                    if image_result['has_text']:
                        image_content.append({
                            'page_num': 1,  # Word文档没有页码概念
                            'image_index': image_index,
                            'extracted_text': image_result['extracted_text'],
                            'type': image_result['type']
                        })
                    image_index += 1
                except Exception as e:
                    print(f"Error processing image in Word document: {e}")
        
        return {
            'text': '\n'.join(text_content),
            'images': image_content
        } 